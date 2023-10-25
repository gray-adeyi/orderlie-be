import json
from abc import ABC
from dataclasses import dataclass, asdict
from enum import Enum
from io import BytesIO
from tempfile import NamedTemporaryFile
from typing import Sequence

from docx import Document
from openpyxl.workbook import Workbook


class FileFormat(str, Enum):
    DOCUMENT = "docx"
    EXCEL = "xlsx"
    PDF = "pdf"
    CSV = "csv"
    JSON = "json"


@dataclass(order=True)
class ExportData:
    headers: Sequence
    rows: Sequence
    metadata: dict | None


class AbstractExporter(ABC):
    def load_data(self, data: ExportData):
        ...

    def export(self) -> BytesIO:
        ...


class JSONExporter(AbstractExporter):
    def load_data(self, data: ExportData):
        self.data = data

    def export(self) -> BytesIO:
        buffer = BytesIO()
        if self.data:
            # TODO: format accordingly
            data = json.dumps(asdict(self.data))
            buffer.write(data.encode("utf8"))
        buffer.seek(0)
        return buffer


class CSVExporter(AbstractExporter):
    def load_data(self, data: ExportData):
        self.data = data

    def export(self) -> BytesIO:
        ...


class PDFExporter(AbstractExporter):
    def load_data(self, data: ExportData):
        self.data = data

    def export(self) -> BytesIO:
        with NamedTemporaryFile() as tmp:
            # canvas = Canvas("hello.pdf")
            # canvas.drawString(72, 72, "Hello, World!")
            # canvas.save()
            tmp.seek(0)
            stream = tmp.read()
            buffer = BytesIO(stream)
            buffer.seek(0)
            return buffer


class XLSXExporter(AbstractExporter):
    def load_data(self, data: ExportData):
        self.data = data
        self.workbook = Workbook()
        self.sheet = self.workbook.active
        self.sheet.title = data.metadata.get("Display Name")

    def export(self) -> BytesIO:
        sheet = self.workbook.active

        # Set the spreadsheet headers
        for col_no, header in enumerate(self.data.headers, start=1):
            self.sheet.cell(row=1, column=col_no).value = header

        # Populate spreadsheet with data
        for row_no, row_value in enumerate(self.data.rows, start=2):
            for col_no, cell_value in enumerate(row_value, start=1):
                self.sheet.cell(row=row_no, column=col_no).value = cell_value

        with NamedTemporaryFile() as tmp:
            self.workbook.save(tmp.name)
            tmp.seek(0)
            stream = tmp.read()
            buffer = BytesIO(stream)
            buffer.seek(0)
            return buffer


class DOCXExporter(AbstractExporter):
    def load_data(self, data: ExportData):
        self.data = data
        self.document = Document()

    def export(self) -> BytesIO:
        self.document.add_heading(self.data.metadata.get("School"), 0)
        self.document.add_paragraph(f"Faculty: {self.data.metadata.get('Faculty')}")
        self.document.add_paragraph(
            f"Department: {self.data.metadata.get('Department')}"
        )
        self.document.add_paragraph(f"Level: {self.data.metadata.get('Level')}")

        self.document.add_paragraph(
            f"Display Name: {self.data.metadata.get('Display Name')}"
        )
        table = self.document.add_table(rows=1, cols=len(self.data.headers))

        # set table headers
        for cell, header in zip(table.rows[0].cells, self.data.headers):
            cell.text = header

        # Populate table with data
        for row_value in self.data.rows:
            row_cells = table.add_row().cells
            for col_no, cell_value in enumerate(row_value):
                row_cells[col_no].text = str(cell_value)

        with NamedTemporaryFile() as tmp:
            self.document.save(tmp.name)
            tmp.seek(0)
            stream = tmp.read()
            buffer = BytesIO(stream)
            buffer.seek(0)
            return buffer


def get_exporter_class(format: FileFormat) -> AbstractExporter:
    return {
        FileFormat.JSON: JSONExporter(),
        FileFormat.CSV: CSVExporter(),
        FileFormat.PDF: PDFExporter(),
        FileFormat.EXCEL: XLSXExporter(),
        FileFormat.DOCUMENT: DOCXExporter(),
    }[format]


def get_media_type(format: FileFormat) -> str:
    return {
        FileFormat.DOCUMENT: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        FileFormat.EXCEL: "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        FileFormat.PDF: "application/pdf",
        FileFormat.CSV: "text/csv",
        FileFormat.JSON: "text/json",
    }[format]
